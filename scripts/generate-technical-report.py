from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Reporte_Tecnico_Mejoras_FarmaPOS.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT = "F2F4F7"
INK = "1F2937"
MUTED = "5B6573"
GREEN = "1F7A5A"  # named FarmaPOS accent override used only on cover/status labels


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tblPr.find(qn(tag))
        if old is not None: tblPr.remove(old)
    tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa"); tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd"); tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa"); tblPr.append(tblInd)
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[idx])); tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); node = OxmlElement("w:tblHeader"); node.set(qn("w:val"), "true"); trPr.append(node)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = str(header)
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd")); cell._tc.tcPr[-1].set(qn("w:fill"), LIGHT)
        for run in cell.paragraphs[0].runs: set_font(run, 9.2, True, DARK_BLUE)
    repeat_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
                for run in p.runs: set_font(run, 8.8, False, INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text); p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.167
    return p


def add_source(doc, label, url):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label}: "); set_font(r, 9, True, MUTED)
    r = p.add_run(url); set_font(r, 9, False, BLUE)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level); p.paragraph_format.keep_with_next = True
    return p


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(1); section.bottom_margin = Inches(1); section.left_margin = Inches(1); section.right_margin = Inches(1)
    section.header_distance = Inches(0.492); section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    tokens = {"Title": (26, DARK_BLUE, 0, 8), "Subtitle": (14, MUTED, 0, 10), "Heading 1": (16, BLUE, 16, 8), "Heading 2": (13, BLUE, 12, 6), "Heading 3": (12, DARK_BLUE, 8, 4)}
    for name, (size, color, before, after) in tokens.items():
        st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(size); st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name != "Subtitle"; st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after); st.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]; st.font.name = "Calibri"; st.font.size = Pt(11); st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.167
    header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("FarmaPOS | Informe técnico XP"); set_font(run, 8.5, False, MUTED)
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Ingeniería de Software I  |  "); set_font(run, 8.5, False, MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); footer._p.append(fld)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(); configure(doc)

    # Editorial cover pattern using the selected business-brief tokens.
    doc.add_paragraph().paragraph_format.space_after = Pt(80)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FARMAPOS"); set_font(r, 13, True, GREEN)
    p.paragraph_format.space_after = Pt(14)
    p = doc.add_paragraph(style="Title"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Reporte técnico de mejoras")
    p = doc.add_paragraph(style="Subtitle"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Operación farmacéutica, atención clínica y análisis gerencial con Extreme Programming")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(36)
    r = p.add_run("Ingeniería de Software I"); set_font(r, 12, True, DARK_BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Fecha: 7 de agosto de 2026"); set_font(r, 10.5, False, MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Autor disponible en el repositorio: Gakilol"); set_font(r, 10.5, False, MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(95)
    r = p.add_run("Next.js · TypeScript · Prisma · PostgreSQL/Neon · Vercel"); set_font(r, 10, True, GREEN)
    doc.add_page_break()

    add_heading(doc, "Resumen ejecutivo", 1)
    p = doc.add_paragraph(); shade_paragraph(p, "EEF6F2"); p.paragraph_format.left_indent = Inches(0.15); p.paragraph_format.right_indent = Inches(0.15); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Resultado: "); set_font(r, 11, True, GREEN)
    p.add_run("se completaron nueve incrementos funcionales integrados a la arquitectura existente: FEFO auditable, recordatorios por rol, caja, receta parcial y segura, resumen temporal del paciente, compras explicables, dashboard gerencial, IA operativa resiliente y reportes ejecutivos optimizados.")
    doc.add_paragraph("El trabajo conservó el modelo serverless y el control de acceso existente. Las operaciones que cambian inventario, receta o caja se ejecutan en transacciones Prisma; las consultas analíticas usan agregaciones, límites y paginación para reducir lecturas y transferencia desde Neon.")
    add_table(doc, ["Área", "Incremento entregado", "Resultado verificable"], [
        ["Farmacia", "FEFO y excepción administrativa", "Lotes vencidos excluidos; trazabilidad DetalleVentaLote; cambio auditado"],
        ["Operación", "Recordatorios y caja", "Alertas internas por rol; apertura, movimientos, cierre y resolución"],
        ["Clínica", "Receta y privacidad", "Ocho estados; surtido parcial; alergias; enlace temporal con hash"],
        ["Gestión", "Compras y KPIs", "Sugerencias explicables; costo de proveedor; agregaciones gerenciales"],
        ["Inteligencia", "IA operativa por rol", "Motor local determinista; Vertex AI opcional; 50 pruebas de permisos y seguridad"],
        ["Analítica", "Reportes ejecutivos", "Una carga inicial; seis KPIs; detalle bajo demanda; exportación completa"],
    ], [1450, 3350, 4560])

    add_heading(doc, "1. Introducción y objetivo", 1)
    doc.add_paragraph("FarmaPOS es un sistema web integrado para farmacia y clínica podológica. El objetivo de esta iteración fue elevar la seguridad operativa, la continuidad del flujo clínico y la capacidad de decisión sin introducir servicios pagados ni procesos persistentes incompatibles con Vercel y Neon en sus modalidades gratuitas.")
    doc.add_paragraph("La prioridad se definió por valor directo: prevenir dispensación insegura, mantener trazabilidad, controlar efectivo, reducir recetas pendientes, facilitar acceso limitado al paciente y convertir datos existentes en decisiones de compra y gestión.")

    add_heading(doc, "2. Diagnóstico inicial", 1)
    add_bullet(doc, "Fortalezas: ventas transaccionales, inventario por lotes, kardex, auditoría, RBAC, clínica SOAP, recetas con cantidades facturadas, reportes y recomendaciones internas.")
    add_bullet(doc, "Brecha FEFO: la deducción ya ordenaba lotes, pero la regla estaba embebida en el endpoint y la interfaz hablaba de FIFO; no existía excepción administrativa explícita.")
    add_bullet(doc, "Brecha operativa: no existía sesión de caja ni conciliación de movimientos y ventas por método de pago.")
    add_bullet(doc, "Brecha clínica: faltaban EN_PREPARACION y LISTA, vínculo relacional Venta-Receta, confirmación de alergias y acceso compartible con aislamiento por paciente.")
    add_bullet(doc, "Brecha analítica: el dashboard inicial cargaba ventas y productos completos al navegador; algunas recomendaciones realizaban consultas por producto.")
    add_bullet(doc, "Riesgo preexistente documentado: los adjuntos de exámenes se escriben en uploads/ mediante fs; ese mecanismo no es almacenamiento durable en Vercel y no fue ampliado.")

    add_heading(doc, "3. Comparación con sistemas referentes", 1)
    doc.add_paragraph("La comparación es funcional, no contractual. Se revisaron páginas oficiales disponibles el 7 de agosto de 2026 para identificar patrones transferibles a un proyecto académico de bajo costo.")
    add_table(doc, ["Referente", "Patrón observado", "Aplicación en FarmaPOS"], [
        ["PioneerRx", "Inventarios múltiples, recomendación de pedidos y detección de artículos sin uso.", "Compras explicables e inventario inmovilizado, sin integración comercial."],
        ["PrimeRx", "Inventario en tiempo real, reposición, costos y fechas de expiración.", "Demanda de 30 días, stock útil, vencimientos y último costo registrado."],
        ["BestRx", "Recetas, inventario, mensajería, reportes, permisos y auditoría integrados.", "Flujo unificado con RBAC y auditoría; mensajería queda sólo diseñada."],
        ["Odoo", "Trazabilidad por lote, fechas de expiración y estrategia FEFO que divide una salida entre lotes.", "Asignación FEFO reutilizable con reparto y selección excepcional auditada."],
        ["Portales clínicos", "Acceso del paciente limitado por identidad, alcance y vigencia.", "Resumen de citas, recetas y resultados autorizados con token aleatorio temporal."],
    ], [1350, 4000, 4010])
    add_source(doc, "PioneerRx", "https://www.pioneerrx.com/pharmacy-software")
    add_source(doc, "PrimeRx", "https://www.primerx.io/pharmacy-inventory-management/")
    add_source(doc, "BestRx", "https://www.bestrx.com/pharmacy-software")
    add_source(doc, "Odoo FEFO", "https://www.odoo.com/documentation/19.0/applications/inventory_and_mrp/inventory/shipping_receiving/removal_strategies/fefo.html")

    add_heading(doc, "4. Requisitos implementados", 1)
    add_heading(doc, "4.1 Funcionales", 2)
    requirements = [
        "RF-01: seleccionar automáticamente el lote vigente con vencimiento más cercano y dividir la cantidad entre lotes cuando sea necesario.",
        "RF-02: permitir sólo a ADMIN cambiar el primer lote, exigir motivo y registrar CAMBIO_LOTE_FEFO.",
        "RF-03: mostrar en el dashboard recordatorios de citas, recetas, vencimientos y stock según el rol.",
        "RF-04: abrir y cerrar caja por usuario; registrar ingreso, retiro y gasto; conciliar efectivo y resolver diferencias por ADMIN.",
        "RF-05: operar recetas en BORRADOR, EMITIDA, EN_PREPARACION, LISTA, USADA_PARCIALMENTE, USADA_COMPLETAMENTE, ANULADA y VENCIDA.",
        "RF-06: impedir cantidades superiores al saldo recetado, asociar Venta con Receta y exigir revisión de alergias.",
        "RF-07: generar un resumen compartible de 24 horas sin exponer otros pacientes ni almacenar el token en claro.",
        "RF-08: recomendar compras con demanda histórica, mínimo, stock actual, vencimiento y costos de proveedores.",
        "RF-09: exponer KPIs gerenciales agregados con rango de fecha y listas limitadas.",
        "RF-10: responder consultas operativas de IA con datos reales, permisos por rol, auditoría y respaldo local cuando el proveedor externo no esté disponible.",
        "RF-11: cargar el resumen de reportes en una solicitud, comparar con el período anterior y obtener los detalles bajo demanda sin degradar las exportaciones.",
    ]
    for item in requirements: add_bullet(doc, item)
    add_heading(doc, "4.2 No funcionales", 2)
    for item in [
        "RNF-01 Seguridad: autenticación existente, RBAC en servidor, tokens SHA-256 y datos clínicos mínimos.",
        "RNF-02 Integridad: transacciones Prisma para venta-receta-inventario, caja y auditoría relacionada.",
        "RNF-03 Concurrencia: bloqueo FOR UPDATE existente en ventas y restricción parcial única para una caja abierta por usuario.",
        "RNF-04 Rendimiento: Promise.all, groupBy, aggregate, SQL agregado, take y paginación; sin polling periódico.",
        "RNF-05 Despliegue: sin Redis, colas, WebSockets, cron externo, almacenamiento nuevo en disco ni APIs pagadas.",
    ]: add_bullet(doc, item)

    add_heading(doc, "5. Historias de usuario y criterios de aceptación", 1)
    stories = [
        ["HU-01 FEFO", "Como vendedor quiero que el sistema sugiera el lote más próximo a vencer para dispensar de forma segura.", "Excluye vencidos y stock cero; reparte entre lotes; conserva DetalleVentaLote; sólo ADMIN cambia lote con motivo auditado."],
        ["HU-02 Recordatorios", "Como usuario quiero ver pendientes relevantes a mi rol al abrir el dashboard.", "DOCTOR ve citas/recetas; EMPLEADO ve stock/vencimientos; ADMIN ve ambos; no hay mensajería externa activa."],
        ["HU-03 Caja", "Como cajero quiero conciliar mi turno y conocer cualquier diferencia.", "Una caja abierta por usuario; movimientos positivos; efectivo esperado correcto; sólo ADMIN resuelve diferencias."],
        ["HU-04 Receta", "Como paciente quiero recibir parcialmente una receta sin perder el saldo.", "No supera pendiente; actualiza estado; venta enlazada; receta vencida o anulada no se surte; alergias requieren confirmación."],
        ["HU-05 Resumen", "Como paciente quiero consultar un resumen limitado mediante un enlace temporal.", "Token aleatorio de 256 bits, hash persistido, expiración y aislamiento por idCliente; sólo resultados autorizados."],
        ["HU-06 Compras", "Como administrador quiero una recomendación explicable para decidir qué comprar.", "Muestra ventas, demanda, stock útil, mínimo, vencimientos, cantidad y proveedor por último costo."],
        ["HU-07 Gerencia", "Como administrador quiero indicadores compactos para detectar riesgos y rentabilidad.", "Ventas, ganancia, bajo stock, vencimientos, top, inmovilizado y caja desde agregaciones del servidor."],
        ["HU-08 IA", "Como usuario quiero consultar inventario y operación aunque el proveedor de IA falle.", "Motor local con intenciones seguras, RBAC en servidor, fuentes visibles y rechazo de diagnóstico o dosificación."],
        ["HU-09 Reportes", "Como administrador quiero analizar el negocio sin esperar múltiples cargas pesadas.", "Resumen único, margen, ticket, variación, detalle perezoso, rango máximo y exportación completa."],
    ]
    add_table(doc, ["Historia", "Necesidad", "Criterios de aceptación"], stories, [1250, 3150, 4960])

    add_heading(doc, "6. Aplicación de Extreme Programming", 1)
    add_heading(doc, "6.1 Planificación e iteraciones", 2)
    iterations = [
        ["1", "FEFO y reglas puras", "Extracción de dominio, prevalidación de stock vigente, selección administrativa y pruebas."],
        ["2", "Recordatorios y caja", "Consultas por rol, modelos aditivos, conciliación y auditoría transaccional."],
        ["3", "Receta y paciente", "Estados, transición, surtido parcial, alergias y resumen temporal protegido."],
        ["4", "Compras y gerencia", "Algoritmo explicable, proveedores, KPIs agregados y componentes del dashboard."],
        ["5", "Integración", "Migración reversible, navegación, tipos, build y documentación."],
        ["6", "Estabilización UI y datos", "Migración aplicada en Neon, historial tolerante a fallos y corrección visual del buscador en modo oscuro."],
        ["7", "IA y reportes", "Motor local con Vertex AI opcional, validación de chat, resumen ejecutivo unificado, RBAC financiero y pruebas de métricas."],
    ]
    add_table(doc, ["Iteración", "Objetivo", "Incremento terminado"], iterations, [900, 2200, 6260])
    add_heading(doc, "6.2 TDD, simplicidad y refactorización", 2)
    doc.add_paragraph("Las reglas con alta densidad de negocio se movieron a funciones puras antes de conectarlas a Prisma. Este diseño permite probar fechas, cantidades y redondeos sin una base de datos, y reduce duplicación entre validación, API y UI.")
    for item in [
        "FEFO: orden por vencimiento, exclusión de vencidos, reparto de cantidad y lote excepcional.",
        "Caja: efectivo esperado = inicial + ventas en efectivo + ingresos - retiros - gastos; diferencia = contado - esperado.",
        "Receta: transiciones permitidas y estado resultante del surtido parcial/completo.",
        "Compras: demanda móvil, stock útil descontando vencimientos y cantidad objetivo.",
        "IA: detección determinista de intenciones operativas, degradación segura y validación de límites del chat.",
        "Reportes: utilidad, margen, ticket promedio, comparación temporal y validación de rangos como funciones puras.",
    ]: add_bullet(doc, item)
    add_heading(doc, "6.3 Retrospectiva técnica", 2)
    doc.add_paragraph("Funcionó bien reutilizar DetalleVentaLote, cantidadFacturada, AuditoriaLog y los roles actuales: el incremento permaneció aditivo. La principal deuda encontrada fue la dispersión de lógica en endpoints extensos; extraer reglas puras redujo riesgo, pero una siguiente iteración debería separar por completo el servicio transaccional de ventas. También se detectó que el comando next lint ya no es válido en Next.js 16 y debe migrarse a ESLint CLI cuando el proyecto incorpore esa dependencia.")

    add_heading(doc, "7. Arquitectura y tecnologías", 1)
    add_table(doc, ["Capa", "Tecnología", "Responsabilidad"], [
        ["Presentación", "Next.js 16 + React 19 + Tailwind", "Dashboard, venta, caja, compras y resumen responsive."],
        ["Aplicación", "Route Handlers + TypeScript + Zod", "Autenticación, permisos, validación e idempotencia."],
        ["Dominio", "Funciones puras TypeScript", "FEFO, caja, receta y recomendaciones verificables."],
        ["Persistencia", "Prisma 6 + PostgreSQL/Neon", "Transacciones, bloqueos, índices, agregaciones y relaciones."],
        ["Despliegue", "Vercel", "Funciones serverless y frontend; sin servidor propio persistente."],
    ], [1450, 2800, 5110])
    add_number(doc, "El navegador solicita únicamente datos del rol autenticado.")
    add_number(doc, "El Route Handler valida sesión, rol y esquema antes de tocar la base.")
    add_number(doc, "La operación crítica se ejecuta en una transacción; la venta bloquea producto/lotes y aplica FEFO.")
    add_number(doc, "Prisma persiste relaciones y auditoría; las respuestas excluyen secretos y hashes.")

    add_heading(doc, "8. Cambios por área", 1)
    add_heading(doc, "8.1 Base de datos", 2)
    doc.add_paragraph("La migración 20260807090000_xp_operacion_integrada es aditiva y contiene rollback. Añade CajaSesion, CajaMovimiento y AccesoPaciente; incorpora Venta.idReceta, Venta.idCaja y ExamenPaciente.autorizadoPortal. Se agregan índices de fecha, estado y relaciones de consulta frecuente, además de un índice único parcial que impide dos cajas ABIERTA para el mismo usuario. La migración fue aplicada en la base Neon configurada y prisma migrate status confirmó que las nueve migraciones están al día.")
    add_heading(doc, "8.2 APIs", 2)
    for item in ["/api/caja", "/api/recordatorios", "/api/dashboard/gerencial", "/api/paciente/acceso", "/api/paciente/resumen/[token]"]: add_bullet(doc, item)
    doc.add_paragraph("Ventas ahora valida stock vigente agregado, alergias, pertenencia de productos a la receta, excepción FEFO y caja abierta. Recetas materializa VENCIDA bajo demanda y expone transiciones controladas.")
    add_heading(doc, "8.3 Interfaz", 2)
    doc.add_paragraph("El dashboard incluye recordatorios e indicadores gerenciales por rol. La venta identifica FEFO y habilita selector de lote sólo a ADMIN. Se añadieron vistas para caja, generación de resumen y consulta pública imprimible. En modo oscuro, los resultados de productos y clientes ahora usan fondo opaco, flujo vertical sin superposición, límites de altura, roles accesibles y bloqueo visible de artículos sin stock.")
    add_heading(doc, "8.4 Seguridad y privacidad", 2)
    add_bullet(doc, "El token del paciente se genera con randomBytes(32), se entrega una vez y se almacena sólo como SHA-256.")
    add_bullet(doc, "La consulta pública parte del idCliente ligado al token; no admite identificadores arbitrarios.")
    add_bullet(doc, "Resultados clínicos requieren autorizadoPortal=true y un enlace creado por ADMIN o DOCTOR.")
    add_bullet(doc, "No se imprimen variables de entorno, expedientes ni tokens en logs.")
    add_heading(doc, "8.5 IA operativa y reportes", 2)
    doc.add_paragraph("El chat valida hasta 20 mensajes y 2,500 caracteres por entrada. Las intenciones operativas conocidas se resuelven primero con herramientas Prisma y formato local determinista, por lo que inventario, vencimientos, ventas y auditoría no dependen de una API externa. Cada herramienta vuelve a validar el rol en servidor y registra el origen del dato. La integración opcional con Vertex AI Express usa el SDK oficial y conserva el motor local ante tiempos de espera o rechazo del proveedor.")
    doc.add_paragraph("Reportes reemplazó nueve solicitudes iniciales por /api/reportes/resumen. El endpoint calcula ventas, compras, COGS, utilidad, margen, ticket promedio, comparación con el período anterior, stock bajo, tendencia y vencimientos en paralelo. Los reportes detallados se cargan sólo al abrir cada pestaña; Excel obtiene todas las hojas explícitamente al exportar. Las rutas financieras y PDF requieren rol ADMIN y limitan el rango a 366 días.")

    add_heading(doc, "9. Optimización para Vercel y Neon", 1)
    add_table(doc, ["Decisión", "Impacto"], [
        ["Recordatorios bajo demanda", "Evita cron, colas y procesos siempre activos; no usa polling."],
        ["Agregaciones en servidor", "Evita descargar tablas completas para KPIs y recomendaciones."],
        ["Límites y paginación", "Caja 20 por página; compras 25; recordatorios 30; top productos 10."],
        ["Índices", "Aceleran estado/fecha, accesos temporales, receta, venta y caja."],
        ["Transacciones cortas", "Mantienen consistencia y reducen ventanas de carrera."],
        ["Resumen ejecutivo unificado", "Reduce la carga inicial de reportes de nueve solicitudes HTTP a una."],
        ["IA local primero", "Responde operaciones frecuentes sin consumo externo y evita latencia cuando la clave avanzada falla."],
    ], [3100, 6260])
    doc.add_paragraph("Neon Free publica 100 CU-horas por proyecto al mes, 0.5 GB por proyecto, escalado a cero tras inactividad y una ventana limitada de restauración. Vercel Hobby publica cuotas de CPU, memoria e invocaciones y está orientado a proyectos personales. Por ello FarmaPOS evita consultas N+1, tareas persistentes y archivos locales como fuente durable.")
    add_source(doc, "Neon Pricing", "https://neon.com/pricing")
    add_source(doc, "Vercel Hobby", "https://vercel.com/docs/plans/hobby")
    add_source(doc, "Vercel Functions", "https://vercel.com/docs/functions")
    add_source(doc, "Archivos en Functions", "https://vercel.com/kb/guide/how-can-i-use-files-in-serverless-functions")

    add_heading(doc, "10. Evidencia de pruebas", 1)
    add_table(doc, ["Verificación", "Resultado", "Evidencia"], [
        ["Pruebas XP de dominio", "APROBADO", "10 reglas: FEFO, vencidos, excepción, caja, receta y compras."],
        ["Suite IA y permisos", "APROBADO", "50 pruebas: RBAC, Zod, prompt injection, límites, FEFO e intenciones locales."],
        ["Suite de reportes", "APROBADO", "8 pruebas: utilidad, margen, ticket, variación, cero y validación de rangos."],
        ["TypeScript", "APROBADO", "npx tsc --noEmit, sin errores."],
        ["Prisma y Neon", "APROBADO", "Esquema válido; migración aplicada; nueve migraciones al día."],
        ["Navegador local", "APROBADO", "IA con Neon, reportes ejecutivos y detalle bajo demanda; cero errores de consola."],
        ["UI modo oscuro", "APROBADO", "Seis KPI legibles, gráfico tematizado, faltante positivo y servicios excluidos del stock."],
        ["Build producción", "APROBADO", "Next.js 16.3.0: 77 rutas, TypeScript incluido y sin advertencias de trazado."],
        ["Lint", "NO EJECUTABLE", "next lint es inválido en Next.js 16; requiere migrar el script a ESLint CLI."],
    ], [2400, 1600, 5360])
    doc.add_paragraph("La advertencia de trazado amplio fue eliminada al acotar los adjuntos clínicos al directorio uploads, validar que las rutas almacenadas no puedan salir de ese límite e indicar al empaquetador cuáles accesos dinámicos están controlados. El build dejó de omitir errores TypeScript y terminó sin advertencias.")

    add_heading(doc, "11. Riesgos, limitaciones y trabajo futuro", 1)
    risks = [
        ["Adjuntos en disco local", "Alto", "Migrar a almacenamiento de objetos cuando exista presupuesto; no ampliar uploads/."],
        ["Resumen, no portal completo", "Medio", "Añadir autenticación del paciente, revocación UI y consentimiento explícito en una fase regulatoria."],
        ["Sin interacciones farmacológicas", "Alto", "Definir interfaz de proveedor clínico validado; no simular reglas médicas incompletas."],
        ["Mensajería externa inactiva", "Bajo", "La estructura de canales existe; activar sólo con consentimiento y proveedor aprobado."],
        ["Cobertura E2E transaccional", "Medio", "Añadir pruebas automatizadas de navegador que creen y reviertan datos aislados en una rama de Neon."],
        ["Lint legado", "Bajo", "Instalar/configurar ESLint compatible con Next.js 16 y sustituir next lint."],
        ["Estados como texto", "Bajo", "Evaluar enum PostgreSQL en una migración posterior tras limpiar valores históricos."],
        ["Clave de IA externa rechazada", "Medio", "La credencial recibida respondió 403 en Vertex y 404 en Gemini clásico; confirmar proveedor, endpoint y permisos. El motor local permanece operativo."],
        ["Dependencia xlsx sin parche", "Medio", "Restringir archivos generados/descargados y evaluar migración a una librería mantenida; npm audit no ofrece corrección."],
    ]
    add_table(doc, ["Riesgo o límite", "Nivel", "Recomendación"], risks, [2550, 1050, 5760])

    add_heading(doc, "12. Conclusión", 1)
    doc.add_paragraph("FarmaPOS avanzó de un conjunto funcional amplio a una operación más segura, explicable y resiliente. El incremento protege la dispensación por lote, enlaza farmacia y clínica, formaliza el efectivo, mantiene una IA operativa sin dependencia externa y entrega reportes comparables con menor carga inicial. La aplicación de XP se evidencia en historias concretas, iteraciones pequeñas, 68 pruebas especializadas, refactorización y una retrospectiva honesta.")
    doc.add_paragraph("La migración ya fue aplicada y verificada en Neon; Caja, IA, reportes y los módulos integrados cargan con datos reales. Antes de publicar en Vercel queda decidir almacenamiento durable para adjuntos, migrar xlsx a una alternativa mantenida, configurar ESLint CLI y obtener una credencial externa con proveedor y endpoint confirmados. El repositorio queda compilable y preparado para la entrega académica.")

    add_heading(doc, "Anexo A. Archivos importantes modificados", 1)
    files = [
        ["prisma/schema.prisma", "Relaciones, caja, acceso paciente, índices."],
        ["prisma/migrations/20260807090000_xp_operacion_integrada/", "Migración SQL y rollback."],
        ["lib/domain/fefo.ts", "Regla FEFO reutilizable."],
        ["lib/domain/cash.ts", "Cálculo de cierre."],
        ["lib/domain/prescriptions.ts", "Estados y transiciones."],
        ["app/api/ventas/route.ts", "Venta, receta, caja y auditoría en transacción."],
        ["app/api/caja/route.ts", "Apertura, movimientos, cierre y resolución."],
        ["app/api/recordatorios/route.ts", "Centro por rol."],
        ["app/api/dashboard/gerencial/route.ts", "KPIs agregados."],
        ["app/api/paciente/", "Generación y consumo del enlace temporal."],
        ["app/ventas/nueva/page.tsx", "FEFO, excepción y alergias."],
        ["app/caja/page.tsx", "Interfaz de caja."],
        ["components/dashboard-reminders.tsx", "Recordatorios en dashboard."],
        ["components/gerencial-panel.tsx", "Indicadores gerenciales."],
        ["lib/ia/local-assistant.ts", "Intenciones y respuestas operativas sin proveedor externo."],
        ["app/api/ia/chat/route.ts", "Vertex AI opcional, validación, permisos y degradación segura."],
        ["app/api/reportes/resumen/route.ts", "KPIs y tendencia ejecutiva en una sola solicitud."],
        ["lib/reportes/metrics.ts", "Cálculos puros de margen, ticket y comparación."],
        ["app/reportes/page.tsx", "Carga diferida, seis KPI, modo oscuro y exportación completa."],
        ["scripts/test-xp-domain.ts", "Pruebas de reglas críticas."],
        ["scripts/test-ia.ts", "50 pruebas de seguridad, permisos y motor local."],
        ["scripts/test-reportes.ts", "8 pruebas de métricas y rangos."],
    ]
    add_table(doc, ["Archivo", "Responsabilidad"], files, [4400, 4960])

    add_heading(doc, "Anexo B. Comandos de validación", 1)
    for cmd in ["npm test", "npm run test:reportes", "npm run test:xp", "npx tsc --noEmit", "npx prisma validate", "npx prisma generate", "npm run build", "npm run lint (detectó script incompatible con Next.js 16)"]: add_bullet(doc, cmd)
    add_heading(doc, "Anexo C. Fuentes consultadas", 1)
    doc.add_paragraph("Fuentes oficiales consultadas el 7 de agosto de 2026. Las capacidades y límites pueden cambiar; deben verificarse antes de una decisión de compra o despliegue.")
    for label, url in [
        ("PioneerRx Pharmacy Software", "https://www.pioneerrx.com/pharmacy-software"),
        ("PrimeRx Inventory Management", "https://www.primerx.io/pharmacy-inventory-management/"),
        ("BestRx Pharmacy Software", "https://www.bestrx.com/pharmacy-software"),
        ("Odoo 19 FEFO", "https://www.odoo.com/documentation/19.0/applications/inventory_and_mrp/inventory/shipping_receiving/removal_strategies/fefo.html"),
        ("Neon Pricing", "https://neon.com/pricing"),
        ("Vercel Hobby", "https://vercel.com/docs/plans/hobby"),
        ("Vercel Functions", "https://vercel.com/docs/functions"),
        ("Vercel file persistence guidance", "https://vercel.com/kb/guide/how-can-i-use-files-in-serverless-functions"),
        ("Vercel payload limit guidance", "https://vercel.com/kb/guide/how-to-bypass-vercel-body-size-limit-serverless-functions"),
        ("Vertex AI Express Mode", "https://docs.cloud.google.com/vertex-ai/generative-ai/docs/samples/googlegenaisdk-vertexai-express-mode"),
    ]: add_source(doc, label, url)

    doc.core_properties.title = "Reporte técnico de mejoras FarmaPOS"
    doc.core_properties.subject = "Ingeniería de Software I - Extreme Programming"
    doc.core_properties.author = "Gakilol"
    doc.core_properties.keywords = "FarmaPOS, XP, Next.js, Prisma, Neon, Vercel, FEFO"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
